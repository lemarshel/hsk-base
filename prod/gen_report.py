"""Generate architecture analysis Word doc -- run: python prod/gen_report.py"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path.home() / "Desktop" / "You are gay.docx"

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(1.0)
sec.left_margin = sec.right_margin = Inches(1.1)

RED   = RGBColor(0xe9, 0x45, 0x60)
NAVY  = RGBColor(0x1a, 0x1a, 0x2e)
GRAY  = RGBColor(0x55, 0x55, 0x55)
LGRAY = RGBColor(0x99, 0x99, 0x99)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RED

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(text)

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(text)

def spacer():
    doc.add_paragraph("")

def bhr(tbl, cells):
    hdr = tbl.rows[0].cells
    for i, v in enumerate(cells):
        hdr[i].text = v
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

def trow(tbl, cells):
    row = tbl.add_row().cells
    for i, v in enumerate(cells):
        row[i].text = v

# =============================================================================
# TITLE PAGE
# =============================================================================
title = doc.add_heading("HSK-Base  /  Lingua Bridge", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = NAVY

sub = doc.add_paragraph("Chinese Module — Codebase Architecture Analysis")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].font.color.rgb = GRAY

dt = doc.add_paragraph("Prepared: March 2026")
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
dt.runs[0].font.color.rgb = LGRAY

doc.add_page_break()

# =============================================================================
# 1. EXECUTIVE SUMMARY
# =============================================================================
h1("1. Executive Summary")
body(
    "HSK-Base is the Chinese language module of the planned Lingua Bridge multi-language "
    "learning platform. Its immediate scope is a searchable 5,364-word Chinese vocabulary "
    "system with TTS, stroke writing practice, topic filtering, snapshot state, export, and "
    "a quiz mode. Its long-term role is to serve as the reference implementation for every "
    "future language module including Arabic and as the foundation for planned Android and iOS apps."
)
spacer()
body(
    "After a significant JS-split refactor completed in March 2026, the codebase is in a "
    "substantially healthier state than it was at the start of the year. The original 2,007-line "
    "hsk.js monolith has been decomposed into 15 focused modules. localStorage keys are now fully "
    "centralized. Critical bugs have been fixed. What remains are structural issues that will "
    "matter as soon as the platform expands beyond the single Chinese vocabulary page, specifically "
    "when the Story Reader, Dashboard, and second language module are added."
)
spacer()

# =============================================================================
# 2. STRENGTHS
# =============================================================================
h1("2. What Is Working Well  (Strengths)")

h2("2.1  Module Split and API Bridge")
body(
    "The original 2,007-line hsk.js monolith has been decomposed into 15 focused modules: "
    "app-config, hsk-api, hsk, tts, ui, palette, lang, sort, filter, storage, export, quiz, "
    "hanzi, render-words, and text-topics. Each module declares what it OWNS, CONSUMES, and "
    "DEPENDS ON in its file header. The window._hsk._register() bridge provides a controlled, "
    "documented API surface between modules. Stub detection via _isStub prevents silent duplicate "
    "registration overwriting real implementations."
)
spacer()

h2("2.2  Single Source of Truth for Word Data")
body(
    "data/words.xlsx is the sole editable master. The Python build pipeline via "
    "prod/make_test_page.py regenerates both index.html and test/index.html from it. Groups, "
    "phonetic data, and POS structure flow through groups.json. No content drift is possible "
    "between files because everything flows from one source."
)
spacer()

h2("2.3  localStorage Key Centralization Now Complete")
body(
    "All 12 localStorage keys now live in window.HSK_LS in app-config.js, covering keys M, F, "
    "L, P, R, V, S, H, PH, LG, PA, and SN. The two early-load preloaders hsk-head.js and "
    "hsk-body.js now also use window.HSK_LS, made possible by moving app-config.js to load "
    "first in the HTML. Renaming any key requires exactly one change in one file."
)
spacer()

h2("2.4  Static Self-Contained Deployment")
body(
    "Both HTML files are fully self-contained at build time with all data inlined. They open "
    "correctly via file:// without a server, which is critical for GitHub Pages deployment and "
    "for future packaging into mobile app webviews."
)
spacer()

h2("2.5  Audit Discipline and Dataset Quality")
body(
    "The project maintains repeatable audit infrastructure under non-usables/audits/, including "
    "integrity checks, merge validation, and tracker verification. The 5,364-word dataset was "
    "built from the strongest available source, then extended by recovering 759 missing words "
    "from official PDFs. This discipline is rare in solo projects and makes the dataset "
    "trustworthy as a curriculum foundation."
)
spacer()

h2("2.6  TTS Resilience")
body(
    "The TTS engine has layered fallback logic including Web Speech API with zh-CN voice "
    "selection, 1.2 second stall detection, Google Translate TTS fallback, and a 15 second outer "
    "failsafe for the browser bug where synthesis starts but the onend event never fires. Stuck "
    "play-button states are prevented by a done-flag deduplication pattern in the Audio chain."
)
spacer()

h2("2.7  Composable Row Visibility Architecture")
body(
    "Filter state is managed entirely through CSS classes on tr elements: hsk-hide, pos-hide, "
    "alpha-hide, sr-hide, and text-hide. This means visibility rules compose correctly across "
    "multiple simultaneous active filters without any JS coordination logic. No filter can "
    "accidentally override another."
)
spacer()

h2("2.8  Flash-Free Theme and Language Preloading")
body(
    "hsk-head.js applies theme colors and palette to the html element before any content is "
    "parsed. hsk-body.js applies language class and column-hide states to the body element "
    "before child elements are parsed. The user sees their saved preferences from the very "
    "first paint with no flash of unstyled content."
)
spacer()

# =============================================================================
# 3. WEAKNESSES
# =============================================================================
h1("3. Problems and Weaknesses")

h2("3.1  CRITICAL  --  The Entire App Is One Giant Pre-Rendered HTML File")
body(
    "The build pipeline inlines 5,364 word rows directly into index.html at build time. "
    "The resulting file is extremely large, several hundred kilobytes of raw DOM, before "
    "the browser has parsed a single script. This causes three concrete problems."
)
bullet("Initial parse time is slow on mobile and low-end devices.")
bullet(
    "The Story Reader, Morpheme Mindmap, Dashboard, and every future page must either be "
    "crammed into the same file or loaded as separate pages with no shared data layer. "
    "The current architecture has no answer for this."
)
bullet(
    "When a second language module is added there is no mechanism to share the app shell. "
    "The entire HTML, CSS, and JS bundle would need to be duplicated for Arabic."
)
body(
    "Root cause: the build pipeline inlines word data as DOM rather than as JSON that a "
    "renderer hydrates at runtime. This made sense when the project was a single standalone "
    "file. It is now actively blocking the platform roadmap."
)
spacer()

h2("3.2  HIGH  --  No Shared Platform Layer Exists Yet")
body(
    "The README and plan.md both describe a lingua-bridge/ root structure with shared/ and "
    "language-specific subdirectories. The codebase does not reflect this. All JS, CSS, and "
    "HTML live flat in hsk-base/. When the Arabic module is built there is no shared/js/, "
    "no shared/css/, and no platform shell to inherit from."
)
bullet("tracker.js is already identified as a shared candidate but has not been separated.")
bullet("The tts.js engine is language-agnostic but is still inside the Chinese module folder.")
bullet(
    "filter.js, sort.js, and storage.js contain no Chinese-specific logic whatsoever. "
    "They are already platform code trapped in language-module files."
)
spacer()

h2("3.3  HIGH  --  Data Pipeline Is Excel-Dependent and Fragile")
body(
    "The only way to update word content is to edit words.xlsx, run make_test_page.py, and "
    "commit the regenerated HTML. This pipeline requires Microsoft Excel or LibreOffice, "
    "requires Python with openpyxl installed, produces no intermediate JSON usable by other "
    "pages, and makes it impossible to update individual words without rebuilding the entire "
    "HTML file. words.json exists as an HTTP-deployment alternative but is not used by the "
    "production page. The runtime reads inlined window.HSK_WORDS, not a fetch call. "
    "The data pipeline and the runtime are structurally disconnected."
)
spacer()

h2("3.4  HIGH  --  Learner State Lives Entirely in localStorage")
body(
    "All progress including learned words, familiar words, snapshots, preferences, row order, "
    "and tracker events is stored in browser localStorage. For a single-device personal tool "
    "this is acceptable. For the platform roadmap it is a hard ceiling."
)
bullet("No cross-device sync is possible.")
bullet("State is wiped on browser data clear or in private-mode sessions.")
bullet("The planned mobile apps have no mechanism to receive state from the web version.")
bullet("The tracker event log has no export path to a backend for future ML or personalization.")
spacer()

h2("3.5  MEDIUM  --  render-words.js Has No Tests")
body(
    "render-words.js builds the entire 5,364-row DOM from window.HSK_WORDS and "
    "window.HSK_GROUPS. Every other JS module depends on this DOM being correct. There are "
    "no unit tests, no integration tests, and no automated verification that the rendered DOM "
    "matches the expected shape documented in architecture.md. A bug in render-words.js "
    "silently breaks everything downstream with no warning."
)
spacer()

h2("3.6  MEDIUM  --  Story Reader, Mindmap, and Dashboard Do Not Exist")
body(
    "The plan.md describes the Story Reader as the heart of the 40-text method and the "
    "Dashboard as the coherence layer that makes the app feel unified. Neither exists in "
    "the codebase. The morpheme mindmap groundwork exists but needs a full visual redesign "
    "away from force-simulation chaos. Until all three exist, the platform is a vocabulary "
    "list, not a curriculum."
)
spacer()

h2("3.7  MEDIUM  --  quiz.js and export.js Have No API Surface")
body(
    "quiz.js and export.js are wrapped in IIFEs and register nothing to window._hsk. "
    "They interact with the DOM directly. There is no mechanism for future pages such as "
    "the Dashboard or Story Reader to trigger a quiz for a specific topic or export filtered "
    "results without duplicating the entire logic."
)
spacer()

h2("3.8  MEDIUM  --  CSS Is a Single Growing Flat File")
body(
    "css/hsk.css is a single flat file at approximately 830 lines containing base resets, "
    "layout rules, vocabulary page styles, theme modes, dark and sepia overrides, phonetic "
    "group controls, quiz styles, export styles, and hanzi popup styles all together. When "
    "the Story Reader and Dashboard pages are added they will either share this file and push "
    "it past 1,500 lines, or duplicate theme variables into separate files causing drift. "
    "There are no CSS custom properties used systematically for spacing or typography "
    "beyond the two palette vars set by hsk-head.js."
)
spacer()

h2("3.9  LOW  --  800 ms Drag Initialization Timing Hack")
body(
    "SortableJS initialization in hsk.js uses setTimeout with an 800 millisecond delay to "
    "wait for the DOM to be fully rendered before attaching drag handles. This is a timing "
    "assumption that will break silently on slow devices and will definitely break when the "
    "vocabulary page transitions to async JSON rendering."
)
spacer()

h2("3.10  LOW  --  No Build Validation or CI")
body(
    "There is no automated check that runs after make_test_page.py to verify the output. "
    "A broken build with malformed HTML, wrong word count, or a missing script tag is "
    "discovered only after deployment to GitHub Pages."
)
spacer()

# =============================================================================
# 4. ARCHITECTURAL RECOMMENDATIONS
# =============================================================================
h1("4. Architectural Recommendations")
body(
    "Ordered by impact on the platform roadmap. Recommendations 1 through 3 are prerequisites "
    "for everything described in plan.md. Recommendations 4 through 6 enable the missing "
    "features. Recommendations 7 and 8 are quality and safety improvements."
)
spacer()

h2("Recommendation 1  --  Switch From Inlined DOM to Runtime JSON Rendering  (CRITICAL)")
body(
    "This is the single most important structural change. Instead of inlining 5,364 rows "
    "of HTML at build time, the build pipeline should output words.json, which already "
    "exists at data/words.json. At runtime, render-words.js fetches this JSON and builds "
    "the DOM once on DOMContentLoaded. The HTML file becomes a small app shell of "
    "approximately 50 KB instead of a 500 KB plus pre-rendered page."
)
body("Concrete benefits of this change:")
bullet("Initial parse time drops dramatically, especially on mobile.")
bullet("The Story Reader, Mindmap, and Dashboard can all read the same words.json.")
bullet("The second language module shares the same shell. Only data/words.json changes.")
bullet("Individual word updates no longer require a full HTML rebuild.")
bullet("The mobile apps can use the same JSON file directly as their data source.")
body(
    "Implementation path: words.json already exists. The required changes are: first, remove "
    "inline word data from make_test_page.py output. Second, have render-words.js call "
    "fetch on DOMContentLoaded. Third, defer hsk.js initialization until render is complete "
    "using a custom DOM event. The file:// offline constraint is already solved by "
    "words-data.js, which can serve as a synchronous fallback."
)
spacer()

h2("Recommendation 2  --  Establish the Lingua Bridge Folder Structure Now  (CRITICAL)")
body(
    "The README and plan.md already define the correct target structure. It should be "
    "created before the Arabic module is started. Migrating the Chinese module now, while "
    "it is the only language module, is the smallest this migration will ever be. "
    "Doing it after Arabic is partially built in the flat structure multiplies the work "
    "and embeds the bad structure into two codebases instead of one."
)
body("Target structure:")
bullet("lingua-bridge/index.html  --  platform landing page")
bullet("lingua-bridge/shared/js/  --  tracker.js, tts-engine.js, filter-core.js, sort-core.js, storage-core.js")
bullet("lingua-bridge/shared/css/  --  base.css, themes.css, components.css")
bullet("lingua-bridge/chinese/  --  vocabulary.html, mindmap.html, stories.html, data/, js/, css/")
bullet("lingua-bridge/arabic/  --  same structure, different data")
body(
    "Files already language-agnostic and ready to move today: tracker.js, filter.js, "
    "sort.js, storage.js, the tts.js engine layer, and the shared constants section "
    "of app-config.js."
)
spacer()

h2("Recommendation 3  --  Build the Story Reader as the Next Feature  (CRITICAL)")
body(
    "The 40-text method is the core of the entire project. Without the Story Reader the "
    "platform is a vocabulary list. Building it next will immediately validate whether the "
    "architecture actually supports the learning method and surface any data or API gaps "
    "before those gaps are inherited by the Arabic module."
)
body("Minimum viable Story Reader based on the requirements in plan.md:")
numbered("One new HTML page: stories.html")
numbered("Sidebar listing all 40 topics with per-topic progress bars from tracker.js")
numbered("Story text area with inline word spans using data-key attribute for hover lookup")
numbered("Hover popup showing pinyin, English meaning, and stroke animation, reusing hanzi.js popup")
numbered("Three visibility toggles: Chinese only, add pinyin, add translation")
numbered("Checklist of 100 target words below each story, persistent via localStorage")
numbered("TTS playback on sentence click, reusing the tts.js engine")
numbered("Previous and next topic navigation")
body(
    "This entire feature can be built without changing the vocabulary page at all. "
    "It is a new HTML file that reads from the same words.json data source."
)
spacer()

h2("Recommendation 4  --  Separate CSS Into a Component System  (HIGH)")
body(
    "Before adding the Story Reader and Dashboard, split hsk.css into at minimum three files."
)
bullet("shared/css/base.css  --  resets, CSS variables with a full token set, typography, layout primitives")
bullet("shared/css/themes.css  --  dark, sepia, and palette override rules")
bullet("chinese/css/vocabulary.css  --  vocabulary table, phonetic groups, quiz, export, hanzi popup")
body(
    "Introduce a systematic CSS variable set for spacing, font sizes, and border radii. "
    "This prevents the Story Reader and Dashboard from duplicating magic numbers copied "
    "from the vocabulary page. Extend the existing --pal-accent and --pal-dark variables "
    "into a full design token set at this point."
)
spacer()

h2("Recommendation 5  --  Register quiz.js and export.js to window._hsk  (MEDIUM)")
body(
    "These two modules currently have no API surface registered to the shared bridge. "
    "Adding two _register calls would expose:"
)
bullet("window._hsk.startQuiz(wordList)  --  so the Dashboard can launch a quiz for any specific topic")
bullet("window._hsk.exportWords(format, filter)  --  so export can be triggered from any page")
body(
    "This is a change of approximately ten lines per file. The payoff is that every future "
    "page gains access to these features without duplicating the logic."
)
spacer()

h2("Recommendation 6  --  Add a Minimal Build Validator  (MEDIUM)")
body(
    "After make_test_page.py runs, a second script at prod/validate_build.py should verify "
    "the output before it is deployed."
)
bullet("Expected word count: 5,364 tr data-key elements must be present")
bullet("All required script tags present and in the correct load order")
bullet("window.HSK_WORDS and window.HSK_GROUPS variable assignments intact")
bullet("At least one row present per POS section")
body(
    "This is approximately 50 lines of Python. No browser, no test framework, no extra "
    "dependencies. It catches build regressions before they reach GitHub Pages."
)
spacer()

h2("Recommendation 7  --  Replace the 800 ms Timeout With an Event-Driven Pattern  (MEDIUM)")
body(
    "The SortableJS initialization in hsk.js currently uses setTimeout with a fixed 800 ms "
    "delay as a proxy for DOM readiness. This timing assumption is fragile and will break "
    "silently on slow devices. The correct fix is to have render-words.js dispatch a custom "
    "event called words-rendered when it finishes building the table. hsk.js then listens "
    "for that event to initialize SortableJS. This removes the timing dependency entirely "
    "and will survive the transition to async JSON rendering in Recommendation 1."
)
spacer()

h2("Recommendation 8  --  Define a Cross-Device State Strategy Before Mobile  (HIGH)")
body(
    "The plan.md explicitly targets Android and iPhone apps. localStorage state cannot cross "
    "the web-to-native device boundary without a defined strategy. Three paths are available."
)
bullet(
    "Path A, simplest: export and import JSON files manually. The learner exports state from "
    "the web app and imports it into the mobile app. No backend required. Works fully offline. "
    "This is the recommended approach for the initial mobile release."
)
bullet(
    "Path B, medium complexity: use a lightweight sync backend such as Supabase or Firebase "
    "free tier. State syncs automatically across all devices once the user creates an account."
)
bullet(
    "Path C, advanced: local-first sync using CRDTs such as Yjs or Automerge. State merges "
    "across devices without any backend. High technical complexity."
)
body(
    "The recommendation is to use Path A for the initial mobile release and plan a clear "
    "upgrade to Path B when the user base grows. The tracker event log export format should "
    "be defined now so it does not need to be redesigned when the mobile app is built."
)
spacer()

# =============================================================================
# 5. PRIORITY MATRIX
# =============================================================================
h1("5. Priority Matrix")
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
bhr(tbl, ["Recommendation", "Impact", "Effort", "Do When"])
for rd in [
    ("1. JSON runtime rendering",         "Critical", "Medium",   "Before Story Reader"),
    ("2. Lingua Bridge folder structure", "Critical", "Low",      "Before Arabic module"),
    ("3. Build Story Reader",             "Critical", "High",     "Next major feature"),
    ("4. Split CSS into components",      "High",     "Low",      "Before Story Reader"),
    ("5. Register quiz/export APIs",      "Medium",   "Very Low", "Next maintenance batch"),
    ("6. Build validator script",         "Medium",   "Very Low", "Next maintenance batch"),
    ("7. Event-driven Sortable init",     "Medium",   "Low",      "With Recommendation 1"),
    ("8. Cross-device state strategy",    "High",     "Low",      "Before mobile planning"),
]:
    trow(tbl, rd)
spacer()
spacer()

# =============================================================================
# 6. WHAT NOT TO DO
# =============================================================================
h1("6. What Not To Do")
bullet(
    "Do not add the Arabic module in the current flat hsk-base/ structure. "
    "Establish the lingua-bridge/ platform layout first or the problem compounds."
)
bullet(
    "Do not build the Dashboard before the Story Reader. The Dashboard aggregates data "
    "from all study tools and needs the Story Reader to exist before it can show meaningful stats."
)
bullet(
    "Do not migrate to a JS framework such as React or Vue. The static file:// deployment "
    "requirement and the planned mobile webview packaging make a framework a liability, not "
    "an asset. Vanilla JS plus a build script is the correct approach for this project."
)
bullet(
    "Do not expand localStorage usage further. Every new feature that stores state should "
    "reuse existing keys from window.HSK_LS rather than introducing new raw string keys. "
    "Every new raw key makes the eventual cross-device sync problem harder to solve."
)
bullet(
    "Do not continue refactoring hsk.js until Recommendation 1 is implemented. The 800 ms "
    "timeout and DOM-scan patterns are symptoms of the inline-DOM architecture, not standalone "
    "bugs. Fixing symptoms before the root cause adds churn without eliminating the problem."
)
spacer()

# =============================================================================
# 7. CURRENT vs TARGET STATE
# =============================================================================
h1("7. Current State vs Target State")
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
bhr(tbl2, ["Area", "Current State", "Target State"])
for rd2 in [
    ("Data delivery",
     "Inlined DOM in HTML, 500 KB plus",
     "fetch(words.json) at runtime, 50 KB shell"),
    ("Folder structure",
     "Flat hsk-base/ monorepo",
     "lingua-bridge/shared/ plus /chinese/"),
    ("Language modules",
     "Chinese only, hardcoded throughout",
     "Chinese and Arabic sharing one platform shell"),
    ("CSS",
     "One flat 830-line file",
     "base.css plus themes.css plus per-page CSS"),
    ("State persistence",
     "localStorage only, no cross-device",
     "localStorage plus JSON export/import for mobile"),
    ("Study features",
     "Vocabulary page only",
     "Vocabulary plus Stories plus Mindmap plus Dashboard"),
    ("Build validation",
     "None, manual inspection only",
     "validate_build.py runs automatically after build"),
    ("Module init timing",
     "800 ms setTimeout assumption",
     "words-rendered custom event, no timing assumption"),
    ("Cross-device sync",
     "Not possible",
     "JSON export/import, Path A, upgradeable to sync"),
]:
    trow(tbl2, rd2)
spacer()
spacer()

# =============================================================================
# 8. CLOSING NOTE
# =============================================================================
h1("8. Closing Note")
body(
    "The codebase is in a genuinely good state for a solo project at this stage of development. "
    "The module split, API bridge, localStorage centralization, data pipeline, and audit "
    "infrastructure are solid foundations that took real effort to establish. The vocabulary "
    "page works correctly and is now maintainable."
)
spacer()
body(
    "The risks identified in this report are architectural, not correctness issues. The current "
    "structure is the right shape for one Chinese vocabulary page. It is not yet the right shape "
    "for a multi-language learning platform with a Story Reader, Dashboard, and mobile apps. "
    "The gap between those two shapes is what the recommendations in Section 4 are meant to "
    "close, in the right order, before the platform expands and the cost of these changes "
    "becomes ten times larger."
)
spacer()
body(
    "If one thing gets done first, it should be switching to runtime JSON rendering "
    "(Recommendation 1). Every other item on the roadmap becomes substantially easier "
    "after that single architectural change is in place."
)

# =============================================================================
# SAVE
# =============================================================================
doc.save(str(OUT))
print("Saved:", OUT)
